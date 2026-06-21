import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_COLOR_INDEX

def generate_report():
    tree = ET.parse('report.xml')
    root = tree.getroot()

    # Create a new Document
    doc = Document()
    doc.add_heading('NexOps Test Execution Report', 0)

    # Collect stats
    testsuite = root.find('.//testsuite')
    if testsuite is not None:
        tests = int(testsuite.attrib.get('tests', 0))
        failures = int(testsuite.attrib.get('failures', 0))
        errors = int(testsuite.attrib.get('errors', 0))
        skipped = int(testsuite.attrib.get('skipped', 0))
        time = testsuite.attrib.get('time', '0')
        passed = tests - failures - errors - skipped
        
        doc.add_heading('Summary', level=1)
        p = doc.add_paragraph()
        p.add_run(f'Total Tests: {tests}\n').bold = True
        p.add_run(f'Passed: {passed}\n').font.color.rgb = None # default, we can just write it
        p.add_run(f'Failures: {failures}\n')
        p.add_run(f'Errors: {errors}\n')
        p.add_run(f'Skipped: {skipped}\n')
        p.add_run(f'Total Duration: {time} seconds\n')

    # Add Test Cases
    doc.add_heading('Test Details', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Test Name'
    hdr_cells[1].text = 'Status'
    hdr_cells[2].text = 'Duration (s)'
    
    for testcase in root.findall('.//testcase'):
        name = testcase.attrib.get('name', 'Unknown')
        classname = testcase.attrib.get('classname', '')
        time_val = testcase.attrib.get('time', '0')
        
        status = 'Passed'
        if testcase.find('failure') is not None:
            status = 'Failed'
        elif testcase.find('error') is not None:
            status = 'Error'
        elif testcase.find('skipped') is not None:
            status = 'Skipped'
            
        row_cells = table.add_row().cells
        row_cells[0].text = f"{classname}.{name}"
        row_cells[1].text = status
        row_cells[2].text = time_val

    doc.save('test_report.docx')
    print("Report saved as test_report.docx")

if __name__ == '__main__':
    generate_report()
